# -*- coding: utf-8 -*-
"""Build UAS SRS MOBILE PRAKTIKUM.docx from the v2.0.0 SRS, adapted to the real app."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), "..")
IMG = os.path.join(ROOT, "docs", "img")
OUT = os.path.join(ROOT, "UAS SRS MOBILE PRAKTIKUM.docx")

PRIMARY = RGBColor(0x00, 0x64, 0x93)
INK = RGBColor(0x20, 0x1C, 0x24)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK

for h, sz in (("Heading 1", 17), ("Heading 2", 14), ("Heading 3", 12)):
    st = doc.styles[h]
    st.font.name = "Calibri"
    st.font.color.rgb = PRIMARY
    st.font.size = Pt(sz)


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def border(paragraph, color="BBBBBB"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "6"); e.set(qn("w:color"), color)
        pbdr.append(e)
    pPr.append(pbdr)


def para(text="", size=11, bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def placeholder(label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    shade(p, "F0F4F8"); border(p, "9DB4C8")
    r = p.add_run("\U0001F4F7  " + label)
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY


def figure(path, caption, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = para(caption, size=9, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = ""
        run = hc[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def fr(code, title, desc, actor, flow=None, status=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{code}: {title}"); r.bold = True; r.font.color.rgb = PRIMARY; r.font.size = Pt(11.5)
    dp = doc.add_paragraph(); dp.paragraph_format.space_after = Pt(2)
    dp.add_run("Deskripsi: ").bold = True; dp.add_run(desc)
    ap = doc.add_paragraph(); ap.paragraph_format.space_after = Pt(2)
    ap.add_run("Aktor: ").bold = True; ap.add_run(actor)
    if flow:
        fp = doc.add_paragraph(); fp.paragraph_format.space_after = Pt(2)
        fp.add_run("Flow:").bold = True
        for i, step in enumerate(flow, 1):
            sp = doc.add_paragraph(style="List Number"); sp.paragraph_format.space_after = Pt(0)
            sp.add_run(step)
    if status:
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)
        sp.add_run("Status implementasi: ").bold = True
        sp.add_run(status).italic = True
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def br(code, title, lines):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{code}: {title}"); r.bold = True; r.font.color.rgb = RGBColor(0x50, 0x60, 0x6E)
    fp = doc.add_paragraph(); fp.paragraph_format.space_after = Pt(2)
    fp.add_run("Fungsi: ").bold = True
    for ln in lines:
        sp = doc.add_paragraph(style="List Bullet"); sp.paragraph_format.space_after = Pt(0)
        sp.add_run(ln)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ============================================================ COVER
for _ in range(3):
    doc.add_paragraph()
para("SOFTWARE REQUIREMENT SPECIFICATION", size=20, bold=True, color=PRIMARY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Aplikasi Mobile Apps — E-Ticketing Helpdesk", size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
figure(os.path.join(IMG, "arch.png"), "Gambar 1. Arsitektur sistem (MVVM, Native Android)", width=5.6)
para("Versi        : 2.0.0", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Tanggal      : 18 Juni 2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Platform     : Native Android (Kotlin + Jetpack Compose)", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("Aplikasi Mobile (Praktikum) — DIV Teknik Informatika, Universitas Airlangga",
     size=11, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================================================ 1. PENDAHULUAN
doc.add_heading("1. Pendahuluan", level=1)
doc.add_heading("1.1. Tujuan Dokumen", level=2)
para("Dokumen ini menjelaskan kebutuhan fungsional dan non-fungsional untuk pengembangan "
     "aplikasi mobile E-Ticketing Helpdesk yang digunakan untuk pelaporan, monitoring, dan "
     "penyelesaian masalah IT atau layanan lainnya.")
para("Fokus dokumen ini adalah:", bold=True, space_after=2)
bullets([
    "Antarmuka pengguna (UI/UX).",
    "Pengelolaan state pada aplikasi mobile.",
    "Interaksi dengan API (backend).",
    "Manajemen data dan penyimpanan file.",
    "Notifikasi dan komunikasi antar pengguna.",
    "Keamanan sistem.",
    "Skalabilitas dan maintainability sistem.",
])

doc.add_heading("1.2. Ruang Lingkup", level=2)
para("Sistem ini mencakup tiga tipe pengguna dengan kemampuan sebagai berikut:")
doc.add_heading("1.2.1. Pengguna", level=3)
bullets(["Membuat tiket keluhan.", "Melihat status tiket.", "Melihat riwayat perjalanan tiket.",
         "Berkomunikasi dengan helpdesk.", "Mendapat notifikasi perubahan tiket.", "Melihat statistik tiket."])
doc.add_heading("1.2.2. Helpdesk", level=3)
bullets(["Membuat tiket.", "Menangani tiket yang ditugaskan.", "Memberikan respon.",
         "Mengubah status tiket.", "Melihat riwayat perjalanan tiket.", "Statistik tiket yang ditugaskan."])
doc.add_heading("1.2.3. Admin", level=3)
bullets(["Membuat tiket.", "Mengelola seluruh tiket.", "Mengelola pengguna.", "Melihat statistik sistem."])

# ============================================================ 2. DESKRIPSI UMUM
doc.add_heading("2. Deskripsi Umum Sistem", level=1)
doc.add_heading("2.1. Perspektif Produk", level=2)
para("Sistem ini merupakan aplikasi client mobile yang dirancang terhubung dengan backend API "
     "(RESTful). Pada versi praktikum ini, aplikasi diimplementasikan secara Native Android dengan "
     "lapisan data dummy in-memory yang merepresentasikan backend, sehingga aplikasi dapat dijalankan "
     "dan dinilai tanpa server. Spesifikasi API dan skema database pada repositori menjadi acuan "
     "integrasi backend nyata tanpa mengubah lapisan domain/UI.")
para("Arsitektur (implementasi nyata):", bold=True, space_after=2)
bullets([
    ("Frontend: ", "Native Android — Kotlin, Jetpack Compose, Material 3."),
    ("Pola arsitektur: ", "MVVM (UI → ViewModel → UseCase → Repository), clean architecture."),
    ("State management: ", "Kotlin Coroutines + StateFlow."),
    ("Navigasi: ", "Navigation Compose (single-activity, NavHost)."),
    ("Data layer: ", "FakeTicketRepository (in-memory) — backend-ready: REST API + MySQL/PostgreSQL."),
])
figure(os.path.join(IMG, "arch.png"), "Gambar 2. Lapisan arsitektur MVVM aplikasi")

doc.add_heading("2.2. Karakteristik Pengguna", level=2)
table(["Tipe User", "Deskripsi"],
      [["Admin", "Pengelola sistem"], ["Helpdesk", "Petugas support"], ["User", "Pelapor tiket"]],
      widths=[1.6, 4.4])

# ============================================================ 3. FUNCTIONAL REQUIREMENT
doc.add_heading("3. Functional Requirement", level=1)
para("Setiap kebutuhan fungsional disertai status implementasi pada aplikasi versi ini. "
     "Legenda: ", space_after=2)
bullets([("Terimplementasi", " — fitur berjalan penuh di aplikasi."),
         ("Representatif", " — alur/UI tersedia memakai data dummy."),
         ("Backend-ready", " — desain & model tersedia, perlu backend/UI lanjutan.")])

doc.add_heading("3.1. Authentikasi & User Management", level=2)
fr("FR-001", "Login", "Pengguna dapat login menggunakan username/email dan password.",
   "Semua tipe pengguna.", status="Terimplementasi (validasi input, pesan ber-tipe AuthMessage).")
fr("FR-002", "Logout", "Pengguna dapat logout dari aplikasi.", "Semua tipe pengguna.",
   status="Terimplementasi (dari layar Profil).")
fr("FR-003", "Register", "Pengguna dapat melakukan pendaftaran akun aplikasi.", "Pengguna.",
   status="Terimplementasi (validasi field, cek duplikasi username/email, konfirmasi password).")
fr("FR-004", "Reset Password", "Pengguna dapat meminta instruksi reset password melalui email.",
   "Semua tipe pengguna.", status="Representatif (verifikasi email terdaftar + konfirmasi; pengiriman email disimulasikan).")
br("BR-001", "Authentication Service", ["Login", "Register", "Logout", "Reset password", "Session management (status login)."])

doc.add_heading("3.2. Management Tiket", level=2)
fr("FR-005", "Pengguna", "Pengguna dapat melakukan permintaan layanan (tiket).", "Pengguna",
   flow=["Membuat tiket.", "Upload laporan (gambar/file: upload atau dari kamera).",
         "Melihat daftar tiket.", "Melihat detail tiket.", "Melihat histori perjalanan tiket.",
         "Mendapatkan notifikasi perubahan tiket.", "Memberikan komentar / reply.", "Melihat statistik tiket."],
   status="Terimplementasi (daftar difilter milik sendiri, detail, histori, komentar, notifikasi, statistik). "
          "Pemilihan sumber lampiran (NONE/CAMERA/FILE) bersifat representatif.")
fr("FR-006", "Helpdesk", "Helpdesk dapat melakukan pengelolaan tiket.", "Helpdesk",
   flow=["Membuat tiket.", "Melihat semua tiket.", "Menangani tiket yang ditugaskan.",
         "Update status pengerjaan tiket.", "Memberikan tanggapan terhadap tiket.",
         "Menutup tiket.", "Melihat statistik tiket."],
   status="Terimplementasi (lihat semua tiket, ubah status, assign, komentar, tutup tiket).")
fr("FR-007", "Admin", "Admin dapat mengelola seluruh tiket dan pengguna.", "Admin",
   flow=["Membuat tiket.", "Melihat semua tiket yang masuk.",
         "Melihat tiket berdasarkan helpdesk yang ditugaskan.", "Menugaskan helpdesk untuk tiket masuk.",
         "Mengubah status tiket.", "Memberikan respon.", "Mengelola daftar pengguna."],
   status="Terimplementasi: lihat semua tiket, assign helpdesk, ubah status, respon. "
          "Manajemen pengguna (daftar & role) backend-ready — model tersedia, UI pengelolaan penuh = pengembangan lanjutan.")
br("BR-002", "Tiket Service", ["Create tiket", "Upload file/image", "Get list tiket", "Get detail tiket",
   "Tracking tiket", "Update status tiket", "Assign tiket", "Delete tiket", "Non-aktifkan pengguna",
   "Menambah komentar", "Menampilkan komentar"])

doc.add_heading("3.3. Notifikasi", level=2)
fr("FR-008", "Notification", "Admin/Helpdesk/Pengguna menerima pemberitahuan terkait aktivitas tiket.",
   "Semua tipe pengguna",
   flow=["Menampilkan pemberitahuan status tiket.", "Navigasi ke halaman terkait."],
   status="Terimplementasi (notifikasi in-app, badge belum-dibaca, tandai semua dibaca).")
br("BR-003", "Notification Service", ["Local Notification / in-app (versi ini).",
   "Backend-ready: Supabase Realtime / Firebase Cloud Messaging (FCM)."])

doc.add_heading("3.4. Dashboard", level=2)
fr("FR-009", "Statistik Tiket", "Menampilkan data ringkasan tiket: Total, Open, Assign, In Progress, Closed.",
   "Semua tipe pengguna", status="Terimplementasi (kartu statistik per status di Dashboard).")
br("BR-004", "Dashboard Service", ["Jumlah tiket.", "Jumlah tiket berdasarkan status tiket."])

doc.add_heading("3.5. Riwayat & Tracking", level=2)
fr("FR-010", "Riwayat Tiket", "Menampilkan riwayat penanganan tiket dari aktivitas semua tipe pengguna.",
   "Semua tipe pengguna", status="Terimplementasi (audit trail aktivitas pada Detail Tiket).")
fr("FR-011", "Tracking Tiket", "Menampilkan status penanganan tiket aktif. User melihat tiket aktif miliknya; "
   "Helpdesk melihat tiket yang ditangani; Admin melihat tiket yang belum ter-close.",
   "Semua tipe pengguna", status="Terimplementasi (status chip + timeline aktivitas pada Detail Tiket).")
br("BR-005", "History Service", ["Menyimpan perubahan status tiket.", "Menyimpan aktivitas pengguna.",
   "Menyediakan tracking tiket."])

# ============================================================ 4. NON-FUNCTIONAL
doc.add_heading("4. Non-Functional Requirement", level=1)
nfr = [
    ("4.1. Performance", ["Lazy loading pada daftar tiket (LazyColumn).", "State reaktif via StateFlow (recompose seperlunya)."]),
    ("4.2. Usability", ["UI responsif.", "Konsisten antar halaman (komponen reusable + tema brand).", "Mudah digunakan."]),
    ("4.3. Compatibility", ["Android (minSdk 24, target/compile SDK 35).", "Mendukung berbagai ukuran layar.",
                            "Catatan: implementasi ini Native Android; dukungan iOS = pengembangan lanjutan (lintas-platform)."]),
    ("4.4. Maintainability", ["Menggunakan clean architecture (pemisahan UI / domain / data).", "MVVM + UseCase + Repository contract."]),
    ("4.5. Security", ["Authentication: login berbasis kredensial (backend-ready JWT / Supabase Auth).",
                       "Authorization: role pengguna (USER, HELPDESK, ADMIN) ditegakkan di ViewModel.",
                       "Build release: R8 (minify), shrink resources, ProGuard, APK ter-signing (keystore)."]),
]
for head, items in nfr:
    doc.add_heading(head, level=2)
    bullets(items)

# ============================================================ 5. UI/UX SCREEN
doc.add_heading("5. UI/UX Screen", level=1)
para("Daftar layar sesuai SRS beserta realisasinya pada aplikasi. Beberapa kebutuhan layar "
     "(Tracking, Setting, Dark/Light) direalisasikan terintegrasi di dalam layar lain.")
table(["Kode", "Layar (SRS)", "Realisasi pada Aplikasi"],
      [
        ["5.1", "Splash Screen", "Tersedia — branding + auto-redirect."],
        ["5.2", "Login Screen", "Tersedia."],
        ["5.3", "Register Screen", "Tersedia."],
        ["5.4", "Forgot Password Screen", "Tersedia (Reset Password)."],
        ["5.5", "Dashboard Screen", "Tersedia — statistik tiket + shortcut."],
        ["5.6", "List Tiket Screen", "Tersedia — difilter per-role."],
        ["5.7", "Detail Tiket Screen", "Tersedia — info, komentar, kontrol status/assign."],
        ["5.8", "Tracking Tiket Screen", "Terintegrasi di Detail Tiket (timeline audit trail)."],
        ["5.9", "Create Tiket Screen", "Tersedia."],
        ["5.10", "Notification Screen", "Tersedia."],
        ["5.11", "Profile Screen", "Tersedia."],
        ["5.12", "Setting Screen", "Terintegrasi di Profil (preferensi & logout)."],
        ["5.13", "Dark & Light Mode", "Tersedia — toggle di Profil, diterapkan global."],
      ], widths=[0.7, 2.5, 3.0])

doc.add_heading("Peta Navigasi", level=2)
figure(os.path.join(IMG, "nav_flow.png"), "Gambar 3. Peta navigasi antar layar")

doc.add_heading("Lampiran Tangkapan Layar (diisi saat pengumpulan)", level=2)
para("Tempelkan screenshot dari emulator/perangkat pada placeholder berikut:", size=10, color=GRAY)
for lbl in ["Splash + Login + Register", "Dashboard (statistik tiket)", "Daftar Tiket + Detail Tiket (timeline)",
            "Create Tiket + Notifikasi", "Profil + Mode Gelap (dark mode)"]:
    placeholder("Screenshot: " + lbl)

# ============================================================ 6. DIAGRAM
doc.add_heading("6. Diagram Sistem", level=1)
doc.add_heading("6.1. Siklus Hidup Tiket", level=2)
figure(os.path.join(IMG, "lifecycle.png"), "Gambar 4. State diagram status tiket", width=5.8)
doc.add_heading("6.2. Entity Relationship Diagram (ERD)", level=2)
figure(os.path.join(IMG, "erd.png"), "Gambar 5. ERD basis data (backend-ready)")
para("DDL CREATE TABLE, penjelasan kolom, dan seed data lengkap tersedia pada "
     "docs/DATABASE_DOCUMENTATION.md.", size=10, italic=True, color=GRAY)
placeholder("Screenshot: tabel database (Structure & Browse) dari phpMyAdmin / MySQL Workbench")

# ============================================================ 7. API
doc.add_heading("7. Antarmuka API (Backend-ready)", level=1)
para("Endpoint REST yang merepresentasikan operasi aplikasi. Detail body, respons, dan skema "
     "data tersedia pada API_DOCUMENTATION.md dan docs/openapi.yaml (dapat diimpor ke Swagger/Postman).")
table(["Method", "Endpoint", "Fungsi"],
      [
        ["POST", "/api/auth/login", "Login pengguna"],
        ["POST", "/api/auth/register", "Registrasi pengguna"],
        ["GET", "/api/tickets", "Daftar tiket (sesuai role)"],
        ["POST", "/api/tickets", "Buat tiket baru"],
        ["PATCH", "/api/tickets/{id}/status", "Ubah status tiket"],
        ["PATCH", "/api/tickets/{id}/assign", "Assign petugas"],
        ["POST", "/api/tickets/{id}/comments", "Tambah komentar"],
        ["GET", "/api/notifications", "Daftar notifikasi"],
      ], widths=[1.0, 2.7, 2.6])
placeholder("Screenshot: pengujian endpoint di Postman / tampilan Swagger UI")

# ============================================================ 8. REFERENSI
doc.add_heading("8. Referensi Dokumen", level=1)
table(["Dokumen", "Isi"],
      [
        ["README.md", "Ikhtisar proyek, tech stack, cara build/run, akun demo"],
        ["docs/FLOW_AND_UIUX.md", "Flow diagram (navigasi, auth, lifecycle, role) + UI/UX tiap layar"],
        ["API_DOCUMENTATION.md", "Endpoint REST API & skema data model"],
        ["docs/openapi.yaml", "Spesifikasi OpenAPI (Swagger/Postman)"],
        ["docs/DATABASE_DOCUMENTATION.md", "Skema relasional, DDL, ERD, seed data"],
        ["docs/VIDEO_TUTORIAL_SCRIPT.md", "Naskah/storyboard video tutorial"],
      ], widths=[2.6, 3.4])

# ============================================================ APPENDIX
doc.add_heading("Lampiran A. Akun Demo", level=1)
para("Password semua akun: 123456", bold=True)
table(["Username", "Role", "Keterangan"],
      [["ahmad / siti / budi", "USER", "Pelapor (lihat & buat tiket sendiri)"],
       ["helpdesk / arif", "HELPDESK", "Menangani semua tiket"],
       ["admin", "ADMIN", "Akses penuh"]], widths=[2.2, 1.4, 2.4])

try:
    doc.save(OUT)
    print("Saved:", os.path.normpath(OUT))
except PermissionError:
    alt = os.path.join(ROOT, "UAS SRS MOBILE PRAKTIKUM (generated).docx")
    doc.save(alt)
    print("LOCKED original; saved to:", os.path.normpath(alt))
print("Sections:", sum(1 for p in doc.paragraphs if p.style.name == "Heading 1"))
